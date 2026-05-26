import fitz  # pymupdf
import docx
import base64
from pathlib import Path

def extract(filepath: str) -> dict:
    path = Path(filepath)
    ext = path.suffix.lower()

    def smart_sample(text: str, head_chars=3000, tail_chars=3000) -> str:
        """Grabs the beginning and end of a document, skipping the middle."""
        if len(text) <= (head_chars + tail_chars):
            return text
        return text[:head_chars] + "\n\n...[CONTENT SKIPPED]...\n\n" + text[-tail_chars:]

    try:
        if ext == ".pdf":
            doc = fitz.open(filepath)
            pages_to_read = min(8, len(doc))  # up to 8 pages, safe for shorter PDFs
            text = "\n".join(doc[i].get_text() for i in range(pages_to_read))
            return {"type": "text", "content": text}

        elif ext == ".docx":
            doc = docx.Document(filepath)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            text = "\n".join(full_text)
            return {"type": "text", "content": smart_sample(text)}

        elif ext == ".doc":
            return {"type": "text", "content": "[Error: Binary .doc files are not supported. Please save/convert as modern .docx format.]"}

        elif ext in (".txt", ".md", ".csv"):
            text = path.read_text(errors="ignore")
            return {"type": "text", "content": smart_sample(text)}

        elif ext in (".png", ".jpg", ".jpeg", ".webp"):
            with open(filepath, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            return {"type": "image", "content": b64}

        else:
            return {"type": "text", "content": f"File: {path.name}"}

    except Exception as e:
        return {"type": "text", "content": f"[Error reading file: {str(e)}]"}